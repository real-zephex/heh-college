from docx import Document

details = [
  ["Prof Sujata V. Mallapur", "Ashwitha", 9490334007, ""],
  ["Dr. N. Jayasudha", "Karampreet", 7658002673, ""],
  ["Prof Shivaganga Patil", "Muskan", 7658091643, "muskancheema077@gmail.com"],
  ["Dr. Prasanth", "Gulshan Nishad", 8948558472, "gulshanjalvanshigolu@gmail.com"],
  ["Dr. Rakesh Mudgal", "Sairam", 8585022461, "unvsr547@gmail.com"],
  ["Dr. Renu Vig", "Ankit Pandey", 9335271760, "officialankitpandey20@gmail.com"],
  ["Prof Bhirud", "Harsh Shukla", 8299585418, ""],
  ["Dr. Piyush Verma", "Bharti Kumari", 6203070486, "bk7701457@gmail.com"]
]

for i in details:
    document = Document()  # Create new document for each iteration
    print(i)
    
    para = document.add_paragraph()
    run = para.add_run("Mail to be sent before the event to the guests\n\n")
    run.bold = True
    run = para.add_run(f"Respected {i[0]} ji,\n")
    run.bold = True

    para.add_run(f"Greetings from Lamrin Tech Skills University Punjab! We hope this mail finds you well. On behalf of the entire LTSU family, it is our pleasure to host you as our distinguished guest for the “First-of-its-kind” 54th ISTE National Annual Convention and Yuva Kaushal Utsav, 2025.\n\n")
    para.add_run(f"This mail is to notify you regarding the assignment of a student protocol officer to assist you during the event. We are pleased to inform you that Mr. {i[1]}, a Bachelor of Technology student, has been designated as your protocol officer. They will personally assist you throughout the event, catering to all your requirements and ensuring a smooth and comfortable experience. They will also serve as your primary point of contact during your visit. We are attaching the protocol officer’s details with the mail.\n")
    para.add_run("\nBest Regards\nVC Conclave Organizing Committee\n")
    run = para.add_run("Student Protocol Details :\n")
    run.bold = True

    para.add_run(f"Name : {i[1]}\n")
    para.add_run(f"Contact no : {i[2]}\n")
    para.add_run(f"Mail Id : {i[3]}\n")

    run = para.add_run("\nIn case of any further assistance, you may contact\n")
    run.bold = True
    
    para.add_run("Student Protocol Incharge:\n")
    para.add_run("1. Meemansa Sharma (8445566794)\n")
    para.add_run("2. Harsh Shukla (8299585418)\n")

    run = para.add_run("\nOverall Heads:\n")
    run.bold = True
    
    para.add_run("1. Dr. HPS Dhami (9876121410)\n")
    para.add_run("2. Dr. Ashutosh Sharma (78892146034)\n")
    
    document.save(f"{i[1]}.docx")